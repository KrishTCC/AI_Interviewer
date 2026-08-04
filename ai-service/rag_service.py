from __future__ import annotations

import glob
import hashlib
import io
import json
import os
import re
from functools import lru_cache
from pathlib import Path
from typing import Any, Optional

from docx import Document
import ollama
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document as LCDocument
from langchain_core.output_parsers import StrOutputParser
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import HumanMessage, SystemMessage
from pypdf import PdfReader

from config import (
    CHROMA_PERSIST_DIR,
    COMPANY_COLLECTION,
    COMPANY_KB_DIR,
    EMBEDDING_MODEL_NAME,
    GEMINI_API_KEY,
    GEMINI_MODEL_NAME,
    KNOWLEDGE_BASE_DIR,
    RAG_TOP_K,
    TECHNICAL_COLLECTION,
    TECHNICAL_KB_DIR,
)


TECH_SKILL_VOCAB = [
    "python",
    "java",
    "javascript",
    "typescript",
    "react",
    "node",
    "express",
    "mongodb",
    "postgresql",
    "mysql",
    "redis",
    "docker",
    "kubernetes",
    "aws",
    "gcp",
    "azure",
    "system design",
    "microservices",
    "rest",
    "graphql",
    "jwt",
    "oauth",
    "ci/cd",
    "linux",
    "tcp",
    "http",
    "os",
    "dbms",
    "dsa",
]


def _read_json_records(directory: str, source_type: str) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    if not os.path.isdir(directory):
        return records

    for file_path in sorted(glob.glob(os.path.join(directory, "*.json"))):
        with open(file_path, "r", encoding="utf-8") as file_handle:
            payload = json.load(file_handle)

        if isinstance(payload, dict):
            payload = [payload]

        for item in payload:
            if not isinstance(item, dict):
                continue

            item = dict(item)
            item["source_type"] = source_type
            item["source_file"] = os.path.basename(file_path)
            records.append(item)

    return records


def _record_to_document(record: dict[str, Any]) -> LCDocument:
    if record.get("source_type") == "company":
        content = (
            f"Company: {record.get('company', '')}\n"
            f"Role: {record.get('role', '')}\n"
            f"Difficulty: {record.get('difficulty', '')}\n"
            f"Topic: {record.get('topic', '')}\n"
            f"Question: {record.get('question', '')}\n"
            f"Ideal Answer: {record.get('ideal_answer', '')}\n"
            f"Follow Up: {record.get('follow_up', '')}\n"
            f"Year: {record.get('year', '')}\n"
            f"Frequency: {record.get('frequency', '')}"
        )
    else:
        tags = ", ".join(record.get("tags", [])) if isinstance(record.get("tags"), list) else record.get("tags", "")
        content = (
            f"Topic: {record.get('topic', '')}\n"
            f"Subtopic: {record.get('subtopic', '')}\n"
            f"Concept: {record.get('concept', '')}\n"
            f"Explanation: {record.get('explanation', '')}\n"
            f"Question: {record.get('question', '')}\n"
            f"Ideal Answer: {record.get('ideal_answer', '')}\n"
            f"Examples: {record.get('examples', '')}\n"
            f"Difficulty: {record.get('difficulty', '')}\n"
            f"Tags: {tags}"
        )

    metadata = {
        "source_type": record.get("source_type", "technical"),
        "source_file": record.get("source_file", ""),
        "company": record.get("company", ""),
        "role": record.get("role", ""),
        "difficulty": record.get("difficulty", ""),
        "topic": record.get("topic", ""),
        "subtopic": record.get("subtopic", ""),
        "tags": ", ".join(record.get("tags", [])) if isinstance(record.get("tags"), list) else record.get("tags", ""),
    }

    return LCDocument(page_content=content, metadata=metadata)


def _knowledge_documents() -> list[LCDocument]:
    company_records = _read_json_records(COMPANY_KB_DIR, "company")
    technical_records = _read_json_records(TECHNICAL_KB_DIR, "technical")
    return [_record_to_document(item) for item in (*company_records, *technical_records)]


@lru_cache(maxsize=1)
def _embedding_model_name() -> str:
    return os.getenv("EMBEDDING_MODEL_NAME", EMBEDDING_MODEL_NAME) or "gemini-embedding-2"


def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name=os.getenv("LOCAL_EMBEDDING_MODEL", _embedding_model_name())
    )


def get_llm(temperature: float = 0.2):
    if GEMINI_API_KEY:
        return ChatGoogleGenerativeAI(
            model=os.getenv("GEMINI_MODEL_NAME", GEMINI_MODEL_NAME),
            google_api_key=GEMINI_API_KEY,
            temperature=temperature,
        )

    raise RuntimeError("GEMINI_API_KEY is required. Set LLM_PROVIDER=gemini and provide a valid Gemini key.")


def _ollama_generate(system_text: str, user_text: str, temperature: float = 0.2) -> str:
    response = ollama.generate(
        model=os.getenv("OLLAMA_MODEL_NAME", "mistral"),
        system=system_text,
        prompt=user_text,
        options={"temperature": temperature},
    )
    return response.get("response", "")


def _invoke_llm(system_text: str, user_text: str, temperature: float = 0.2) -> str:
    if GEMINI_API_KEY:
        try:
            llm = ChatGoogleGenerativeAI(
                model=os.getenv("GEMINI_MODEL_NAME", GEMINI_MODEL_NAME),
                google_api_key=GEMINI_API_KEY,
                temperature=temperature,
            )
            response = llm.invoke([SystemMessage(content=system_text), HumanMessage(content=user_text)])
            return response.content if hasattr(response, "content") else str(response)
        except Exception as exc:
            message = str(exc).lower()
            if "quota" not in message and "resource_exhausted" not in message and "rate" not in message:
                raise

    try:
        return _ollama_generate(system_text, user_text, temperature=temperature)
    except Exception:
        return ""


@lru_cache(maxsize=1)
def _vectorstores_ready() -> tuple[Chroma, Chroma]:
    embeddings = get_embeddings()
    persist_dir = CHROMA_PERSIST_DIR
    os.makedirs(persist_dir, exist_ok=True)

    company_store = Chroma(
        collection_name=COMPANY_COLLECTION,
        persist_directory=persist_dir,
        embedding_function=embeddings,
    )
    technical_store = Chroma(
        collection_name=TECHNICAL_COLLECTION,
        persist_directory=persist_dir,
        embedding_function=embeddings,
    )

    if company_store._collection.count() == 0 or technical_store._collection.count() == 0:
        docs = _knowledge_documents()
        company_docs = [doc for doc in docs if doc.metadata.get("source_type") == "company"]
        technical_docs = [doc for doc in docs if doc.metadata.get("source_type") == "technical"]

        if company_store._collection.count() == 0 and company_docs:
            company_store.add_documents(company_docs)
        if technical_store._collection.count() == 0 and technical_docs:
            technical_store.add_documents(technical_docs)

    return company_store, technical_store


def _normalize_json_text(raw_text: str) -> str:
    cleaned = raw_text.strip()
    cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()
    return cleaned


def _safe_json_loads(raw_text: str, default: Any) -> Any:
    cleaned = _normalize_json_text(raw_text)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(0))
            except json.JSONDecodeError:
                return default
        return default


def _format_docs(docs: list[LCDocument]) -> str:
    blocks: list[str] = []
    for index, doc in enumerate(docs, start=1):
        metadata = doc.metadata or {}
        block = (
            f"Context {index}:\n"
            f"Source Type: {metadata.get('source_type', '')}\n"
            f"Company: {metadata.get('company', '')}\n"
            f"Role: {metadata.get('role', '')}\n"
            f"Topic: {metadata.get('topic', '')}\n"
            f"Tags: {metadata.get('tags', '')}\n"
            f"Content: {doc.page_content}"
        )
        blocks.append(block)
    return "\n\n".join(blocks)


def _store_for_mode(mode: str) -> Chroma:
    company_store, technical_store = _vectorstores_ready()
    return company_store if mode == "company" else technical_store


def _resume_text_to_profile(resume_text: str) -> dict[str, Any]:
    system_text = (
        "You are a resume intelligence engine. Return only JSON with keys: "
        "skills (array of strings), technologies (array of strings), projects (array of strings), "
        "experience_summary (string), target_topics (array of strings), inferred_role (string), "
        "strengths (array of strings). Do not add markdown or commentary."
    )
    user_text = f"Resume text:\n{resume_text}\n\nExtract a concise structured profile."
    response_text = _invoke_llm(system_text, user_text, temperature=0.1)
    parsed = _safe_json_loads(response_text, {})
    if isinstance(parsed, dict) and parsed:
        return parsed
    return _heuristic_resume_profile(resume_text)


def _heuristic_resume_profile(resume_text: str) -> dict[str, Any]:
    lower_text = resume_text.lower()
    found_skills = sorted({skill for skill in TECH_SKILL_VOCAB if skill in lower_text})
    lines = [line.strip(" -•\t") for line in resume_text.splitlines() if line.strip()]
    project_lines = [line for line in lines if any(token in line.lower() for token in ["project", "built", "developed", "implemented"])]
    role = "Software Engineer"

    if any(token in lower_text for token in ["frontend", "react", "ui"]):
        role = "Frontend Developer"
    elif any(token in lower_text for token in ["backend", "api", "microservice"]):
        role = "Backend Engineer"
    elif any(token in lower_text for token in ["data science", "machine learning", "ml"]):
        role = "Machine Learning Engineer"

    target_topics: list[str] = []
    if any(skill in ["python", "javascript", "java"] for skill in found_skills):
        target_topics.append("coding fundamentals")
    if any(skill in ["mongodb", "postgresql", "mysql", "redis"] for skill in found_skills):
        target_topics.append("databases")
    if any(skill in ["docker", "kubernetes", "aws", "ci/cd"] for skill in found_skills):
        target_topics.append("system design")
    if any(skill in ["react", "javascript", "typescript"] for skill in found_skills):
        target_topics.append("frontend development")

    return {
        "skills": found_skills,
        "technologies": found_skills,
        "projects": project_lines[:5],
        "experience_summary": " ".join(lines[:8])[:1200],
        "target_topics": target_topics or ["software engineering"],
        "inferred_role": role,
        "strengths": found_skills[:8],
    }


def parse_resume_file(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename or "resume.txt").suffix.lower()

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if suffix in {".docx", ".doc"}:
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    return file_bytes.decode("utf-8", errors="ignore")


def _candidate_queries_from_profile(profile: dict[str, Any]) -> list[str]:
    queries: list[str] = []
    queries.extend(profile.get("skills", []))
    queries.extend(profile.get("technologies", []))
    queries.extend(profile.get("target_topics", []))
    queries.extend(profile.get("projects", [])[:3] if isinstance(profile.get("projects"), list) else [])
    queries.extend([profile.get("inferred_role", "")])
    return [query for query in queries if isinstance(query, str) and query.strip()]


def _semantic_similarity(text_a: str, text_b: str) -> float:
    if not text_a.strip() or not text_b.strip():
        return 0.0

    embeddings = get_embeddings()
    vector_a = embeddings.embed_query(text_a)
    vector_b = embeddings.embed_query(text_b)
    if not vector_a or not vector_b:
        return 0.0

    numerator = sum(a * b for a, b in zip(vector_a, vector_b))
    denom_a = sum(a * a for a in vector_a) ** 0.5
    denom_b = sum(b * b for b in vector_b) ** 0.5
    if denom_a == 0 or denom_b == 0:
        return 0.0
    return max(0.0, min(1.0, numerator / (denom_a * denom_b)))


def _extract_missing_concepts(answer_text: str, ideal_answer: str, concept_tags: list[str]) -> list[str]:
    answer_lower = answer_text.lower()
    missing: list[str] = []

    for tag in concept_tags:
        if tag and tag.lower() not in answer_lower:
            missing.append(tag)

    if missing:
        return missing[:6]

    ideal_keywords = [token.strip() for token in re.split(r"[,;/()\n]", ideal_answer) if len(token.strip()) > 4]
    for keyword in ideal_keywords:
        if keyword.lower() not in answer_lower and keyword not in missing:
            missing.append(keyword)
        if len(missing) >= 6:
            break

    return missing


def _heuristic_evaluation(
    *,
    question: str,
    question_type: str,
    role: str,
    level: str,
    user_answer: str,
    user_code: str,
    ideal_answer: str,
    concept_tags: list[str],
    semantic_similarity: float,
) -> dict[str, Any]:
    answer_text = f"{user_answer}\n{user_code}".lower()
    missing_concepts = _extract_missing_concepts(answer_text, ideal_answer, concept_tags)
    concepts_correctly_explained = [tag for tag in concept_tags if tag and tag.lower() in answer_text]

    technical_score = round((semantic_similarity * 70) + max(0, 30 - len(missing_concepts) * 5))
    if question_type == "coding" and not user_code.strip():
        technical_score = min(technical_score, 35)

    confidence_score = min(100, max(0, int(len(user_answer.strip()) * 1.5)))
    overall_score = round((technical_score + confidence_score) / 2)

    feedback_parts = [
        f"This answer was evaluated heuristically for {role} at {level} level.",
        f"The response covered {len(concepts_correctly_explained)} core concepts.",
    ]
    if missing_concepts:
        feedback_parts.append(f"Missing concepts: {', '.join(missing_concepts[:5])}.")
    else:
        feedback_parts.append("The answer aligned well with the reference concepts.")

    return {
        "technicalScore": int(max(0, min(100, technical_score))),
        "confidenceScore": int(max(0, min(100, confidence_score))),
        "overallScore": int(max(0, min(100, overall_score))),
        "semanticSimilarity": round(semantic_similarity, 3),
        "aiFeedback": " ".join(feedback_parts),
        "idealAnswer": ideal_answer or "No reference answer available.",
        "conceptsCorrectlyExplained": concepts_correctly_explained,
        "missingConcepts": missing_concepts,
    }


def retrieve_context(
    *,
    mode: str,
    company: Optional[str] = None,
    role: Optional[str] = None,
    difficulty: Optional[str] = None,
    topic: Optional[str] = None,
    resume_profile: Optional[dict[str, Any]] = None,
    count: int = RAG_TOP_K,
) -> list[LCDocument]:
    store = _store_for_mode("company" if mode == "company" else "technical")

    if mode == "company":
        query_parts = [company or "", role or "", difficulty or "", topic or "company interview"]
        query = " ".join(part for part in query_parts if part).strip()
        filter_kwargs: dict[str, Any] = {}
        if company:
            filter_kwargs["company"] = company
        if role:
            filter_kwargs["role"] = role
        if difficulty:
            filter_kwargs["difficulty"] = difficulty
        if topic:
            filter_kwargs["topic"] = topic
        return store.similarity_search(query or f"{company} {role} interview", k=count, filter=filter_kwargs or None)

    if resume_profile:
        queries = _candidate_queries_from_profile(resume_profile)
        query = " ".join(queries[:6]) if queries else resume_profile.get("experience_summary", "resume interview")
        return store.similarity_search(query, k=count)

    query = " ".join(part for part in [role, difficulty, topic, "technical interview"] if part)
    return store.similarity_search(query or "general technical interview", k=count)


def _build_question_generation_prompt(mode: str) -> ChatPromptTemplate:
    if mode == "company":
        system_text = (
            "You are a senior interview designer creating company-specific interview questions grounded in retrieved evidence. "
            "Return only JSON with key 'questions'. Each question item must include questionText, questionType, idealAnswer, followUp, topic, conceptTags, sourceContext. "
            "Do not add markdown or commentary."
        )
    elif mode == "resume":
        system_text = (
            "You are a resume-aware interview generator. Use the extracted candidate profile and retrieved technical context to create personalized interview questions. "
            "Return only JSON with key 'questions'. Each question item must include questionText, questionType, idealAnswer, followUp, topic, conceptTags, sourceContext. "
            "Do not add markdown or commentary."
        )
    else:
        system_text = (
            "You are a technical interview generator. Use the retrieved technical knowledge to produce realistic interview questions. "
            "Return only JSON with key 'questions'. Each question item must include questionText, questionType, idealAnswer, followUp, topic, conceptTags, sourceContext. "
            "Do not add markdown or commentary."
        )

    return ChatPromptTemplate.from_messages(
        [
            ("system", system_text),
            (
                "human",
                "Mode: {mode}\nCompany: {company}\nRole: {role}\nDifficulty: {difficulty}\nLevel: {level}\nInterview Type: {interview_type}\nCount: {count}\nTopic: {topic}\nResume Profile: {resume_profile}\nRetrieved Context:\n{context}\n\nGenerate exactly {count} interview questions grounded in the context.",
            ),
        ]
    )


def _fallback_questions_from_context(docs: list[LCDocument], count: int, interview_type: str) -> list[dict[str, Any]]:
    questions: list[dict[str, Any]] = []
    coding_cutoff = int(count * 0.2) if interview_type == "coding-mix" else 0

    for index, doc in enumerate(docs[:count]):
        metadata = doc.metadata or {}
        is_coding = index < coding_cutoff
        base_question = metadata.get("question") or re.split(r"Question:\s*", doc.page_content, flags=re.IGNORECASE)[-1].strip().splitlines()[0]
        ideal_answer = metadata.get("ideal_answer") or re.split(r"Ideal Answer:\s*", doc.page_content, flags=re.IGNORECASE)[-1].strip()
        topic = metadata.get("topic") or metadata.get("subtopic") or "general"
        concept_tags = [tag.strip() for tag in str(metadata.get("tags", "")).split(",") if tag.strip()]

        questions.append(
            {
                "questionText": base_question,
                "questionType": "coding" if is_coding else "oral",
                "idealAnswer": ideal_answer,
                "followUp": metadata.get("follow_up", ""),
                "topic": topic,
                "conceptTags": concept_tags,
                "sourceContext": doc.page_content,
            }
        )

    while len(questions) < count:
        questions.append(
            {
                "questionText": "Explain a key technical concept from your experience.",
                "questionType": "oral",
                "idealAnswer": "A strong answer explains the concept clearly, includes trade-offs, and uses a real example.",
                "followUp": "How would you apply this concept in production?",
                "topic": "technical depth",
                "conceptTags": ["technical depth"],
                "sourceContext": "fallback",
            }
        )

    return questions[:count]


def generate_questions(
    *,
    mode: str,
    company: Optional[str],
    role: str,
    difficulty: str,
    level: str,
    interview_type: str,
    count: int,
    topic: Optional[str] = None,
    resume_text: Optional[str] = None,
    resume_filename: Optional[str] = None,
) -> dict[str, Any]:
    resume_profile = _resume_text_to_profile(resume_text) if resume_text else None
    docs = retrieve_context(
        mode=mode,
        company=company,
        role=role,
        difficulty=difficulty,
        topic=topic,
        resume_profile=resume_profile,
        count=max(RAG_TOP_K, count),
    )
    context = _format_docs(docs)

    if mode == "company":
        system_text = (
            "You are a senior interview designer creating company-specific interview questions grounded in retrieved evidence. "
            "Return only JSON with key 'questions'. Each question item must include questionText, questionType, idealAnswer, followUp, topic, conceptTags, sourceContext. "
            "Do not add markdown or commentary."
        )
    elif mode == "resume":
        system_text = (
            "You are a resume-aware interview generator. Use the extracted candidate profile and retrieved technical context to create personalized interview questions. "
            "Return only JSON with key 'questions'. Each question item must include questionText, questionType, idealAnswer, followUp, topic, conceptTags, sourceContext. "
            "Do not add markdown or commentary."
        )
    else:
        system_text = (
            "You are a technical interview generator. Use the retrieved technical knowledge to produce realistic interview questions. "
            "Return only JSON with key 'questions'. Each question item must include questionText, questionType, idealAnswer, followUp, topic, conceptTags, sourceContext. "
            "Do not add markdown or commentary."
        )

    user_text = (
        f"Mode: {mode}\nCompany: {company or ''}\nRole: {role}\nDifficulty: {difficulty}\nLevel: {level}\n"
        f"Interview Type: {interview_type}\nCount: {count}\nTopic: {topic or ''}\n"
        f"Resume Profile: {json.dumps(resume_profile or {}, ensure_ascii=False)}\nRetrieved Context:\n{context}\n\n"
        f"Generate exactly {count} interview questions grounded in the context."
    )
    response_text = _invoke_llm(system_text, user_text, temperature=0.35)

    parsed = _safe_json_loads(response_text, {})
    questions = parsed.get("questions") if isinstance(parsed, dict) else None
    if not isinstance(questions, list) or not questions:
        questions = _fallback_questions_from_context(docs, count, interview_type)

    normalized_questions: list[dict[str, Any]] = []
    for index, item in enumerate(questions[:count]):
        if isinstance(item, str):
            normalized_questions.append(
                {
                    "questionText": item,
                    "questionType": "coding" if interview_type == "coding-mix" and index < int(count * 0.2) else "oral",
                    "idealAnswer": "",
                    "followUp": "",
                    "topic": topic or "technical interview",
                    "conceptTags": [],
                    "sourceContext": "llm-string-fallback",
                }
            )
            continue

        normalized_questions.append(
            {
                "questionText": item.get("questionText") or item.get("question") or item.get("prompt") or "Explain a technical concept.",
                "questionType": item.get("questionType") or ("coding" if interview_type == "coding-mix" and index < int(count * 0.2) else "oral"),
                "idealAnswer": item.get("idealAnswer") or item.get("ideal_answer") or "",
                "followUp": item.get("followUp") or item.get("follow_up") or "",
                "topic": item.get("topic") or topic or "technical interview",
                "conceptTags": item.get("conceptTags") or item.get("concept_tags") or [],
                "sourceContext": item.get("sourceContext") or item.get("source_context") or context,
            }
        )

    return {
        "questions": normalized_questions,
        "mode": mode,
        "company": company,
        "role": role,
        "difficulty": difficulty,
        "level": level,
        "resume_profile": resume_profile,
        "retrieved_context": context,
        "model_used": os.getenv("GEMINI_MODEL_NAME", GEMINI_MODEL_NAME) if GEMINI_API_KEY else os.getenv("OLLAMA_MODEL_NAME", "mistral"),
    }


def evaluate_answer(
    *,
    question: str,
    question_type: str,
    role: str,
    level: str,
    user_answer: str = "",
    user_code: str = "",
    ideal_answer: str = "",
    concept_tags: Optional[list[str]] = None,
) -> dict[str, Any]:
    concept_tags = concept_tags or []
    answer_text = f"{user_answer}\n{user_code}".strip()
    semantic_similarity = _semantic_similarity(answer_text, ideal_answer or question)
    missing_concepts = _extract_missing_concepts(answer_text, ideal_answer, concept_tags)

    system_text = (
        "You are a strict interview evaluator. Return only JSON with keys: technicalScore, confidenceScore, aiFeedback, idealAnswer, conceptsCorrectlyExplained, missingConcepts. "
        "Use the provided semantic similarity as a signal, but do not copy it directly. Evaluate technical accuracy, completeness, and clarity. "
        "Do not add markdown or commentary."
    )
    user_text = (
        f"Role: {role}\nLevel: {level}\nQuestion: {question}\nQuestion Type: {question_type}\n"
        f"Ideal Answer: {ideal_answer or 'No reference answer available.'}\nConcept Tags: {', '.join(concept_tags)}\n"
        f"Semantic Similarity: {round(semantic_similarity, 3)}\nUser Answer: {user_answer or 'No verbal answer provided.'}\n"
        f"User Code: {user_code or 'No code provided.'}\n\nProvide feedback with concrete missing concepts and an interview-grade score."
    )
    response_text = _invoke_llm(system_text, user_text, temperature=0.15)

    if not response_text.strip():
        return _heuristic_evaluation(
            question=question,
            question_type=question_type,
            role=role,
            level=level,
            user_answer=user_answer,
            user_code=user_code,
            ideal_answer=ideal_answer,
            concept_tags=concept_tags,
            semantic_similarity=semantic_similarity,
        )

    parsed = _safe_json_loads(response_text, {})
    technical_score = int(parsed.get("technicalScore", 0)) if isinstance(parsed, dict) else 0
    confidence_score = int(parsed.get("confidenceScore", 0)) if isinstance(parsed, dict) else 0
    ai_feedback = parsed.get("aiFeedback", "") if isinstance(parsed, dict) else ""
    ideal_answer_out = parsed.get("idealAnswer", ideal_answer) if isinstance(parsed, dict) else ideal_answer
    concepts_correctly_explained = parsed.get("conceptsCorrectlyExplained", []) if isinstance(parsed, dict) else []

    if not isinstance(concepts_correctly_explained, list):
        concepts_correctly_explained = []

    llm_technical = max(0, min(100, technical_score))
    llm_confidence = max(0, min(100, confidence_score))
    semantic_component = round(semantic_similarity * 100)
    blended_technical = round((llm_technical * 0.6) + (semantic_component * 0.4))
    blended_confidence = round((llm_confidence * 0.7) + min(100, len(answer_text) * 2)) if answer_text else llm_confidence
    overall_score = round((blended_technical + blended_confidence) / 2)

    return {
        "technicalScore": int(blended_technical),
        "confidenceScore": int(min(100, blended_confidence)),
        "overallScore": int(overall_score),
        "semanticSimilarity": round(semantic_similarity, 3),
        "aiFeedback": ai_feedback or "The answer has been evaluated against the reference answer and retrieved concepts.",
        "idealAnswer": ideal_answer_out or ideal_answer,
        "conceptsCorrectlyExplained": concepts_correctly_explained,
        "missingConcepts": parsed.get("missingConcepts", missing_concepts) if isinstance(parsed, dict) else missing_concepts,
    }
